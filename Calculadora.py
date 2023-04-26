
import pandas as pd
import os
df=pd.read_excel("DatosHotelEC.xlsx", index_col=0)

if not os.path.exists("images"):
    os.mkdir("images")
    
import warnings
warnings.filterwarnings("ignore")
#diccionario de variables
diccionario_variables={
    
'ID':'IDENTIFICACIÓN DEL ALOJAMIENTO',
'YEAR': 'Año',
'SEM': 'Semanas de apertura al año',
'D1':'PROVINCIA',
'D2':'NÚMERO DE ESTRELLAS',
'D3':'TIPOLOGÍA DE ALOJAMIENTO (por categoría)',
'D4':'TIPOLOGÍA DE ALOJAMIENTO (por ubicación)',
'D5':'FUNCIONAMIENTO',
'D6':'ESCALA',
'D7':'ESTACIÓN DEL AÑO EN FUNCIONAMIENTO',
'D8':'PROPIEDAD',
'D9':'SUPERFICIE DEL ALOJAMIENTO (m2)',
'D10':'Nº DE HABITACIONES (Uds.)',
'D11':'Nº PERNOCTACIONES POR AÑO (huésped-noche/año)',
'D12':'Nº HABITACIONES OCUPADAS POR AÑO (habitaciones ocupadas/año)',
'D13':'SUPERFICIE DE HABITACIÓN ESTÁNDAR (m2)',
    
'P1B':'Fracción de la inversión buenas prácticas circulares (% sobre total)',
'P2B':'Plantilla que ha recibido formación relacionada con la circularidad (%)',
'P3B':'Proveedores que operan con un código de conducta circular (%)',
'PBP1':'Comunicaciones o guías relacionadas con circularidad para los clientes',
'PBP2': 'Principios básicos de la EC',
'PBP3': 'Estrategia o Plan de circularidad',
'PBP4': 'Indicadores relacionados con EC',
'PBP5': 'Informe de sostenibilidad',                      
'PBP6': 'Selección de proveedores de proximidad o certificación de sostenibilidad',
'PBP7': 'Compra de productos reacondicionables, reciclables,...',
'PBP8': 'Alojamiento con certificado de sostenibilidad',
'PBP9':'Premio buenas prácticas de sostenibilidad y/o de EC',
'PBP10':'Colaboración con stakeholders sobre sostenibilidad y/o circularidad',
'PBP11':'Implantación tecnologías digitales para sostenibilidad y/o circularidad',
'PBP12':'Personal encargado EC, gestión de residuos, sostenibilidad, etc.',
'PBP13':'Los trabajadores conocen qué es la EC y sus principios',
'PBP14':'Concienciación a la plantilla en materia de sostenibilidad y/o EC',
'PBP15':'Planes de formación para sus trabajadores en materias de EC',
'PBP16':'Servicios asociados de movilidad sostenible',
'PBP17':'Productos y/o servicios turísticos con menor huella de carbono',
'PBP18':'Conocimiento de la legislación sobre sostenibilidad y EC',
'PBP19':'Conocimiento programas de financiación pública sostenibilidad y EC',

###ENERGÍA
'E1B': 'Huella de carbono anual (Teq CO2 año)',
'E2B': 'Certificación energética del edificio',
'E3B': 'Capacidad de autoabastecimiento de energía (%)',
'E4B': 'Potencia renovable instalada (kWh)',
'E5B': 'Capacidad de almacenamiento (kWh)',
'E6': 'Consumo anual total de energía del alojamiento (MWh/año)',
'E7':'Consumo anual de energía térmica del alojamiento (MWh/año)',
'E8':'Consumo anual de energía eléctrica del alojamiento (MWh/año)',
'E9':'Calificación energética sistemas de iluminación',
'E10':'Calificación energética instalaciones de cocina',
'E11':'Calificación energética instalaciones de habitaciones',
'E12':'Calificación anergética ascensores',
'EBP1': 'Programas o iniciativas para compensar la huella de carbono',
'E13':'Huella de carbono compensada (KWh/año)',
'EBP21':'Sensibilización y fomento en clientes reducción consumo de energía',
'EBP22': 'Sensibilización y fomento en trabajadores reducción consumo de energía',
'EBP3':'Planes de mantenimiento de las instalaciones de gestión de energía', 
'EBP4':'Sistema de gestión de energía en habitaciones',
'EBP5':'Actuación para la rehabilitación energética de los edificios',
'EBP61':'Actuación de mejora aislamiento edificio',
'EBP62':'Actuación de mejora aislamiento instalaciones y equipamiento',
'EBP7':'Protecciones solares exteriores',

    
    

###AGUA
    
'A1B':'Capacidad de autoabastecimiento de agua (%)',
'A2B':'Consumo anual de agua red pública (l/año)',
  
'ABP1':'Fomento reducción de agua entre los clientes',
'ABP2':'Plan de mantenimiento instalaciones de gestión del agua',
'ABP3':'Instalaciones de tratamiento, filtrado y reaprovechamiento de aguas',
'ABP41':'Prácticas circulares en habitaciones',
'ABP42':'Prácticas circulares otras instalaciones del alojamiento',
'ABP5':'Especies vegetales autóctonas que requieren menor consumo de agua',
    
###RESIDUOS    

'R1B':'Recogida selectiva de residuos (m3/noche)',
'R2B':'Reciclaje de residuos de obras, reformas y demoliciones (%)',
'R1':'Residuos orgánicos (m3/año)',
'R2':'Residuos de plástico (m3/año)',
'R3':'Residuos de papel y cartón (m3/año)',
'R4':'Residuos de vidrio (m3/año)',
'R5':'Otros residuos recogidos selectivamente que pueden ser valorizados Aceites, tejidos (m3/año)',
'R6': 'Residuos no depositados selectivamente (fraccción resto) m3/año',
'RBP1':'Plan de prevención y gestión de residuos',
'RBP2':'Prácticas circulares en materia de gestión de residuos (no alimentación)',
'RBP3':'Fomento en clientes de separación de residuos',
'RBP41':'Empleo de productos químicos sostenibles',
'RBP42':'Empleo de productos textiles sostenibles',
'RBP43':'Empleo de productos de aseo y cosmética sostenibles',
'RBP44':'Empleo de equipamiento y mobiliario sostenible',
'RBP45':'Empleo de otros productos sostenibles',
'RBP5': 'Actuación para eliminar, sustituir o minimizar los envases y embalajes',
'RBP6': 'Actuación para envases y embalajes reutilizables, reciclables o compostables',
'RBP7': 'Sistema de depósito, devolución y retorno de los envases',
'RBP8': 'Uso de materiales de construcción revalorizados o sostenibles',
'RBP9': 'Tratamiento de los residuos de construcción',




#DESPERDICIO ALIMENTARIO
    
'RA1B': 'Consumo de productos de kilómetro cero (% del total)',
'RA2B': 'Cesta de la compra que minimiza el uso de envases (%)',
'RABP1': 'Seguimiento para conocer alimentos no consumidos en el alojamiento',
'RABP11': 'Fracción de alimentos no consumidos (%)',
'RABP2':  'Fomento de la reducción del desperdicio alimentario entre clientes', 
'RABP3':  'Prácticas circulares optimización del consumo y reducción del desperdicio de alimentos',
'RABP4':  'Uso de bioalimentos, alimentos de origen orgánico y alimentos ecológicos y sostenibles',
'RABP5':  'Prácticas circulares de gestión de residuos de alimentación y otros residuos orgánicos',
'RABP6':  'Uso de excedentes de alimentos para otros usuarios',


#CÁLCULOS planificación y gestión circular
#Buenas prácticas

'PBP'  : 'Subdimensión Planificación Actuaciones',

    
#CÁLCULOS ENERGÍA

#Certificación energética
'E2B_val':'Certificación Edificio',
'EI':'Certificación instalaciones',
#Buenas prácticas
'EBP':'Subdimensión Energía Actuaciones',

#cálculos por huesped, habitación y m2
'E1B1':'Huella de carbono, por huesped y noche (kg CO2 eq. / huesped y noche)',    
'E1B2':'Huella de carbono, por habitación ocupada y noche (kg CO2 eq. / hab. Ocupada y noche)',
'E1B3':'Huella de carbono, por m2',

'E3B1':'Consumo anual del alojamiento de energía procedente de energías renovables (MWh/ huesped y noche)',
'E3B2':'Consumo anual del alojamiento de energía procedente de energías renovables (MWh/ hab. Ocupada y noche)',
'E3B3':'Consumo anual del alojamiento de energía procedente de energías renovables (MWh/por m2)',

'E61':'Consumo de energía del alojamiento por huesped y noche (Kwh/huesped y noche)',
'E71': 'Consumo de energía térmica del alojamiento por huesped y noche (Kwh/huesped y noche)',
'E81': 'Consumo de energía eléctrica del alojamiento por huesped y noche (Kwh/huesped y noche)',

'E62':  'Consumo de energía del alojamiento por habitación ocupada y noche (Kwh/hab Ocupada y noche)',
'E72':  'Consumo de energía térmica del alojamiento por habitación ocupada y noche (Kwh/hab Ocupada y noche)',
'E82':  'Consumo de energía eléctrica del alojamiento por habitación ocupada y noche (Kwh/hab Ocupada y noche)',

'E63':  'Consumo de energía del alojamiento por m2 (Kwh/m2)',
'E73':  'Consumo de energía térmica del alojamiento por m2 (Kwh/m2)',
'E83':  'Consumo de energía eléctrica del alojamiento por m2 (Kwh/m2)', 

    
#CÁLCULOS AGUA
'A1B1':'Consumo de agua de autoabastecimiento /huesped (l/año)',
'A1B2':'Consumo de agua de autoabastecimiento /habitación ocupada (l/año)',
'A1B3':'Consumo de agua de autoabastecimiento /m2 (l/año)',

'A2B1':'Consumo de agua red /huesped (l/año)',
'A2B2':'Consumo de agua red pública/habitación ocupada (l/año)',
'A2B3':'Consumo de agua red pública/m2 (l/año)',
    

    
'ABP':'Subdimensión Agua Actuaciones' ,

#CÁLCULOS RESIDUOS
'R1P':'Residuos orgánicos (kg/año)',
'R2P':'Residuos de plástico (kg/año)',
'R3P':'Residuos de papel y cartón (kg/año)',
'R4P':'Residuos de vidrio (kg/año)',
'R5P':'Otros residuos (kg/año)' ,
'R6P': 'Residuos no depositados selectivamente (fraccción resto) kg/año',
'R1F':'Fracción de residuos orgánicos (%)',
'R2F':'Fracción de residuos de plástico (%)',
'R3F':'Fracción de residuos de papely cartón (%)',
'R4F':'Fracción de residuos de vidrio (%)',
'R5F':'Fracción de otros residuos valorizados (%)' ,
'R6F':'Fracción resto (%)' ,
'R1FP':'Fracción en peso de residuos orgánicos (%)',
'R2FP':'Fracción en peso de residuos de plástico (%)',
'R3FP':'Fracción en peso de residuos de papely cartón (%)',
'R4FP':'Fracción en peso de residuos de vidrio (%)',
'R5FP':'Fracción en peso de otros residuos valorizados (%)' ,
'R6FP':'Fracción en peso resto (%)' , 
'RVal':'% de residuos recogidos selectivamente frente al total (en volumen)',


'R11':'Residuos orgánicos (l/huesped)',
'R21':'Residuos de plástico (l/huesped)',
'R31':'Residuos de papel y cartón (l/huesped)',
'R41':'Residuos de vidrio (l/huesped)',
'R51':'Otros residuos valorizados (l/huesped)',
'R61':'Residuos no valorizados (l/huesped)',
    
'R12':'Residuos orgánicos (l/por habitación ocupada)',
'R22':'Residuos de plástico (l/por habitación ocupada)',
'R32':'Residuos de papel y cartón (l/por habitación ocupada)',
'R42':'Residuos de vidrio (l/por habitación ocupada)',
'R52':'Otros residuos valorizados (l/por habitación ocupada)',
'R62':'Residuos no valorizados (l/por habitación ocupada)',
      
'R13':'Residuos orgánicos (l/m2)',
'R23':'Residuos de plástico (l/m2)',
'R33':'Residuos de papel y cartón (l/m2)',
'R43':'Residuos de vidrio (l/m2)',
'R53':'Otros residuos valorizados (l/m2)',
'R63':'Residuos no valorizados (l/m2)',
    
'RTot': 'Total de residuos (m3)',
'RTotP': 'Total de residuos en peso (kg)',

'RBP':'Subdimensión Residuos Actuaciones' ,

#CÁLCULOS DESPERDICIO ALIMENTARIO

'RABP':'Subdimensión Residuos Alimentarios Actuaciones' ,


} 


# In[3]:


df_scale=pd.read_excel('HerrEC_scale_new.xlsx', index_col=0)
#funciones

# Establecimiento de la valoración según la escala
def escala(df,indicador,valor_indicador):
    if valor_indicador<df.loc[indicador]['Bajo']:
        escala_indicador='Bajo'
    elif valor_indicador<df.loc[indicador]['Medio']:
        escala_indicador='Medio'
    elif valor_indicador<df.loc[indicador]['Alto']:
        escala_indicador='Alto'
    else:
        escala_indicador='Muy Alto'
    return (escala_indicador)

# Establecimiento de la valoración según la escala certificación energética
def escala_certificacion(indicador):
    if indicador=='G':
        valor_indicador=0
    elif indicador=='F':
        valor_indicador=1/6
    elif indicador=='E':
        valor_indicador=2/6
    elif indicador=='D': 
        valor_indicador=3/6
    elif indicador=='C':
        valor_indicador=4/6
    elif indicador=='B':
        valor_indicador=5/6
    elif indicador=='A':
        valor_indicador=1
    return(valor_indicador)

# Establecimiento de la valoración según la escala certificación energética
def escala_certificacion_inversa(indicador,i):
    valor=df.iloc[i][indicador]
    if valor==0:
        valor_indicador='G'
    elif valor<1/6:
        valor_indicador='F'
    elif valor<2/6:
        valor_indicador='E'
    elif valor<3/6: 
        valor_indicador='D'
    elif valor<4/6:
        valor_indicador='C'
    elif valor<5/6:
        valor_indicador='B'
    else: valor_indicador='A'
    return(valor_indicador)


#cambiar pregunta huella de carbono total E1B

#CÁLCULOS
#CÁLCULO DIMENSIONES

#####PLANIFICACIÓN

#PLANIFICACIÓN Inversión P1B



def calculos (df,i):
    
    
    
######CÁLCULOS DIMENSIÓN PLANIFICACIÓN   
    #INVERSIÓN PB1
    #incorporar a tabla de resultados
    #P1B_punt=Planificacion*Plan_Inversion*P1B[i]
    

    #PLANIFICACIÓN FORMACIÓN P2_3B
    #df.P2_3B[i] = (df.P2B[i]+df.P3B[i])/2

    #PLANIFICACIÓN ACTUACIONES PBP
    df.PBP[i]= df.PBP1[i]+df.PBP2[i]+df.PBP3[i]+df.PBP4[i]+df.PBP5[i]+ df.PBP6[i]+df.PBP7[i]+df.PBP8[i]+df.PBP9[i]+df.PBP10[i]+df.PBP11[i]+df.PBP12[i]+df.PBP13[i]+df.PBP14[i]+df.PBP15[i]+df.PBP16[i]+df.PBP17[i]+df.PBP18[i]+df.PBP19[i]

    
#####ENERGÍA
    #ENERGÍA Certificación enegética Edificio EB2
    df.E2B_val[i]=escala_certificacion(df.E2B[i])
    
    #ENERGÍA Certificación enegética Instalaciones
    df.EI[i]=(escala_certificacion(df.E9[i])+escala_certificacion(df.E10[i])+escala_certificacion(df.E11[i])+escala_certificacion(df.E12[i]))/4



    # ENERGIA Parámetros
    #Huella de carbono (T CO2 eq. año) E1B

    
    #Huella de carbono, por huesped en kgeq
    df.E1B1[i]=df.E1B[i]*1000/df.D11[i]
    #Huella de carbono, por habitación ocupada 
    df.E1B2[i]=df.E1B[i]*1000/df.D12[i]
    #Huella de carbono, por m2
    df.E1B3[i]=df.E1B[i]*1000/df.D9[i]
    

    #'E3B': 'Capacidad de autoabastecimiento de energía (%)' ,
    #'E4B': 'Potencia renovable instalada (kWh)',
    #'E5B': 'Capacidad de almacenamiento (kWh)' ,
    
    
    #Consumo de energía huésped y noche (kWh/ huésped y noche)
    df.E61[i]=df.E6[i]*1000/df.D11[i]#total
    df.E71[i]=df.E7[i]*1000/df.D11[i]#térmica
    df.E81[i]=df.E8[i]*1000/df.D11[i]#eléctrica
    df.E3B1[i]=(df.E3B[i]*df.E6[i]/100)*1000/df.D11[i]# renovable total 
   
    
    
    
    # Consumo de energía  por habitación ocupada y noche
    df.E62[i]=df.E6[i]*1000/df.D12[i]
    df.E72[i]=df.E7[i]*1000/df.D12[i]
    df.E82[i]=df.E8[i]*1000/df.D12[i]
    df.E3B2[i]=(df.E3B[i]*df.E6[i]/100)*1000/df.D12[i]
    
    # Consumo de energía  por m2
    df.E63[i]=df.E6[i]*1000/df.D9[i]
    df.E73[i]=df.E7[i]*1000/df.D9[i]
    df.E83[i]=df.E8[i]*1000/df.D9[i]
    df.E3B3[i]=(df.E3B[i]*df.E6[i]/100)*1000/df.D9[i]
    
    #ENERGÍA actuaciones
    df.EBP[i]= df.EBP1[i]+df.EBP21[i]+df.EBP22[i]+df.EBP3[i]+df.EBP4[i]+ df.EBP5[i]+df.EBP61[i]+df.EBP62[i]+df.EBP7[i]


    #####AGUA
    #Capacidad de autoabastecimiento A1B
    
    df.A1B1[i]=(df.A1B[i]*df.A2B[i]/100)/df.D11[i] #l por huesped autoabastecimiento 
    df.A1B2[i]=(df.A1B[i]*df.A2B[i]/100)/df.D12[i] #l por habitación autoabastecimiento 
    df.A1B3[i]=(df.A1B[i]*df.A2B[i]/100)/df.D9[i] #l por m2
    
    #Consumo por habitación ocupada /año A2B1 #Consumo anual de agua, por huésped (L/huésped y año) (m3/huésped y año)
    df.A2B1[i]=df.A2B[i]/df.D11[i] #l por huesped
    df.A2B2[i]=df.A2B[i]/df.D12[i] #l por habitación
    df.A2B3[i]=df.A2B[i]/df.D9[i] #l por m2

    
    #D11=Nº PERNOCTACIONES POR AÑO (huésped-noche/año)',
    #valor mínimo de consumo 290, valor entre intervalos 50
    #Consumo Agua
    #df.A1_2B[i]=((df.A1B[i]/100)+df.A2B1_val[i])/2


    #AGUA actuaciones
    df.ABP[i]= df.ABP1[i]+df.ABP2[i]+df.ABP3[i]+df.ABP41[i]+df.ABP42[i]+ df.ABP5[i]


######RESIDUOS

    
    df.RTot[i]=df.R1[i]+ df.R2[i]+df.R3[i]+df.R4[i]+df.R5[i]+df.R6[i]
    df.RTotP[i]=df.R1P[i]+ df.R2P[i]+df.R3P[i]+df.R4P[i]+df.R5P[i]+df.R6P[i]
    df.RVal[i]=(df.RTot[i]-df.R6[i])*100/df.RTot[i]
    df.R1F[i]=df.R1[i]*100/df.RTot[i]
    df.R2F[i]=df.R2[i]*100/df.RTot[i]
    df.R3F[i]=df.R3[i]*100/df.RTot[i]
    df.R4F[i]=df.R4[i]*100/df.RTot[i]
    df.R5F[i]=df.R5[i]*100/df.RTot[i]
    df.R6F[i]=df.R6[i]*100/df.RTot[i]

    df.RTotP[i]=df.R1P[i]+ df.R2P[i]+df.R3P[i]+df.R4P[i]+df.R5P[i]+df.R6P[i]
    df.R1FP[i]=df.R1P[i]*100/df.RTotP[i]
    df.R2FP[i]=df.R2P[i]*100/df.RTotP[i]
    df.R3FP[i]=df.R3P[i]*100/df.RTotP[i]
    df.R4FP[i]=df.R4P[i]*100/df.RTotP[i]
    df.R5FP[i]=df.R5P[i]*100/df.RTotP[i]
    df.R6FP[i]=df.R6P[i]*100/df.RTotP[i]
    
    df.R11[i]=df.R1[i]*1000/df.D11[i]
    df.R21[i]=df.R2[i]*1000/df.D11[i]
    df.R31[i]=df.R3[i]*1000/df.D11[i]
    df.R41[i]=df.R4[i]*1000/df.D11[i]
    df.R51[i]=df.R5[i]*1000/df.D11[i]
    df.R61[i]=df.R6[i]*1000/df.D11[i]
    
    df.R12[i]=df.R1[i]*1000/df.D12[i]
    df.R22[i]=df.R2[i]*1000/df.D12[i]
    df.R32[i]=df.R3[i]*1000/df.D12[i]
    df.R42[i]=df.R4[i]*1000/df.D12[i]
    df.R52[i]=df.R5[i]*1000/df.D12[i]
    df.R62[i]=df.R6[i]*1000/df.D12[i]
    
    df.R13[i]=df.R1[i]*1000/df.D9[i]
    df.R23[i]=df.R2[i]*1000/df.D9[i]
    df.R33[i]=df.R3[i]*1000/df.D9[i]
    df.R43[i]=df.R4[i]*1000/df.D9[i]
    df.R53[i]=df.R5[i]*1000/df.D9[i]
    df.R63[i]=df.R6[i]*1000/df.D9[i]
    
    #RESIDUOS actuaciones
    df.RBP[i]=  df.RBP1[i]+df.RBP2[i]+df.RBP3[i]+df.RBP41[i]+df.RBP42[i]+df.RBP43[i]+df.RBP44[i]+df.RBP45[i]+ df.RBP5[i]+ df.RBP6[i]+ df.RBP7[i]+ df.RBP8[i]++ df.RBP9[i]

#####DESPERDICIO ALIMENTARIO

    #Alimentos km0 RA1B
    #Alimentos reducción envases RA2B
    #ALIMENTOS actuaciones
    df.RABP[i]=  df.RABP1[i]+df.RABP2[i]+df.RABP3[i]+df.RABP4[i]+ df.RABP5[i]+df.RABP6[i]
    
    
    return df
df=calculos(df,df.index[0]) 
df=calculos(df,df.index[1]) 
#df.to_excel("DatosHotelEC_relleno.xlsx")

indicadores_dim=['P1B','P2B','P3B','PBP','E1B','E2B_val','EI','E3B','E61','E63','EBP','A1B','A2B2','A2B3','ABP','RVal','R2B','RBP','RA1B','RA2B','RABP']

i=0
def Evaluacion_indicador (indicador,Value):
    #Value =valor del indicador
    V_dim=df_scale.Pdim[indicador]#puntos de la dimensión
    V_sdim=df_scale.Psdim[indicador]#%asignada a la subdimensión

    n_min=df_scale.Min[indicador] #lee el valor mínimo de la hoja de escalas
    n_bajo=df_scale.Bajo[indicador] #lee el valor bajo de la hoja de escalas
    n_medio=df_scale.Medio[indicador]#lee el valor Medio de la hoja de escalas
    n_alto=df_scale.Alto[indicador] #lee el valor Alto de la hoja de escalas
    n_max=df_scale.Max[indicador] #lee el valor máximo de la hoja de escalas
    
    if df_scale.Dir[indicador]==0:
        # Establecimiento de la valoración según la escala
        if Value<n_min:
            Valoracion='Bajo'
            Puntos =0
        elif Value<n_bajo:
            Valoracion='Bajo'
            Puntos= round(((Value-n_min)/(n_max-n_min))*(V_dim*V_sdim),2)
        elif Value<n_medio:
            Valoracion='Medio'
            Puntos= round(((Value-n_min)/(n_max-n_min))*(V_dim*V_sdim),2)
        elif Value<n_alto:
            Valoracion='Alto'
            Puntos= round(((Value-n_min)/(n_max-n_min))*(V_dim*V_sdim),2)
        elif Value<n_max:
            Valoracion='Muy Alto'
            Puntos= round(((Value-n_min)/(n_max-n_min))*(V_dim*V_sdim),2)
        else:
            Valoracion='Muy Alto'
            Puntos=round((V_dim*V_sdim),2)
    
    elif df_scale.Dir[indicador]==1:
            # Establecimiento de la valoración según la escala dirección opuesta (valores bajos son positivos para el medio ambiente)
        if Value<n_max:
            Valoracion='Muy Alto'
            Puntos =round((V_dim*V_sdim),2)
        elif Value<n_alto:
            Valoracion='Muy Alto'
            Puntos= round(((n_min-Value)/(n_min-n_max))*(V_dim*V_sdim),2)
        elif Value<n_medio:
            Valoracion='Alto'
            Puntos= round(((n_min-Value)/(n_min-n_max))*(V_dim*V_sdim),2)
        elif Value<n_bajo:
            Valoracion='Medio'
            Puntos= round(((n_min-Value)/(n_min-n_max))*(V_dim*V_sdim),2)
        elif Value<n_min:
            Valoracion='Bajo'
            Puntos= round(((n_min-Value)/(n_min-n_max))*(V_dim*V_sdim),2)
        else:
            Valoracion='Bajo'
            Puntos=0
    return Valoracion,Puntos    


#CALCULO DE LA DATAFRAME DE RESULTADOS
columnas_results=['Indicador','Dimensión','Subdimensión', 'Valor', 'Puntuación asignada','Nivel','Valor_y1','Puntuación asignada_y1','Nivel_y1']
results= pd.DataFrame(columns=columnas_results)
for indicador in indicadores_dim:
    Valoracion=Evaluacion_indicador(indicador,df.iloc[i][indicador])[0]
    Puntos=Evaluacion_indicador(indicador,df.iloc[i][indicador])[1]
    Valoracion_y1=Evaluacion_indicador(indicador,df.iloc[i+1][indicador])[0]
    Puntos_y1=Evaluacion_indicador(indicador,df.iloc[i+1][indicador])[1]
    hotel_results={'Indicador': indicador,'Dimensión': df_scale.Dim[indicador],'Subdimensión':df_scale.Sdim[indicador],
                   'Valor':df.iloc[i][indicador],'Puntuación asignada':Puntos,'Nivel':Valoracion, 
                   'Valor_y1':df.iloc[i+1][indicador],'Puntuación asignada_y1':Puntos_y1,'Nivel_y1':Valoracion_y1}
    results.loc[indicador] = pd.Series(hotel_results)


    
#agregación de valores de subdimensión PARA CONFORMAR LAS DIMENSIONES
num_subdim=len(results)
dimensiones=['Planificación y Gestión Circular','Energía','Agua','Residuos','Desperdicio Alimentario']
suma_puntos_totales=0 #set a cero los puntos totales
suma_puntos_totales_y1=0
for dimension in dimensiones:
    suma_puntos_dim=0#set cero los puntos para cada dimensión una vez empieza a iterar
    suma_puntos_dim_y1=0#set cero los puntos para cada dimensión una vez empieza a iterar
    for x in range(0,num_subdim):
        if results.iloc[x]['Dimensión']==dimension:
            suma_puntos_dim=suma_puntos_dim+results.iloc[x]['Puntuación asignada']
            suma_puntos_dim_y1=suma_puntos_dim_y1+results.iloc[x]['Puntuación asignada_y1']
    Valoracion=Evaluacion_indicador(dimension,suma_puntos_dim)[0]  #va a la función Evaluación_indicador y extrae la valoración
    Puntos=Evaluacion_indicador(dimension,suma_puntos_dim)[1] #va a la función  Evaluación_indicador  y extrae la valoración
    Valoracion_y1=Evaluacion_indicador(dimension,suma_puntos_dim_y1)[0]  #va a la función Evaluación_indicador y extrae la valoración
    Puntos_y1=Evaluacion_indicador(dimension,suma_puntos_dim_y1)[1] #va a la función  Evaluación_indicador  y extrae la valoración
    hotel_results={'Indicador': dimension,'Dimensión': "",'Subdimensión':"", 
                   'Valor':suma_puntos_dim,'Puntuación asignada':Puntos,'Nivel':Valoracion,
                  'Valor_y1':suma_puntos_dim_y1,'Puntuación asignada_y1':Puntos_y1,'Nivel_y1':Valoracion_y1}#creamos un diccionario para añadir una nueva fila a resultados
    results.loc[dimension] = pd.Series(hotel_results)#añadimos los datos como una nueva fila a la tabla resultados
    suma_puntos_totales= suma_puntos_totales+Puntos#variable para calcular lo puntos totales que va acumulando los puntos de las dimensiones
    suma_puntos_totales_y1= suma_puntos_totales_y1+Puntos_y1

#agregación de valores de PARA CONFORMAR EL NIVEL DE CIRCULARIDAD
hotel_results={'Indicador': 'Nivel de Circularidad','Dimensión': '','Subdimensión':'', 
               'Valor':suma_puntos_totales,'Puntuación asignada':Evaluacion_indicador('Nivel de Circularidad',suma_puntos_totales)[1] ,'Nivel':Evaluacion_indicador('Nivel de Circularidad',suma_puntos_totales)[0],
              'Valor_y1':suma_puntos_totales_y1,'Puntuación asignada_y1':Evaluacion_indicador('Nivel de Circularidad',suma_puntos_totales_y1)[1],'Nivel_y1':Evaluacion_indicador('Nivel de Circularidad',suma_puntos_totales_y1)[0]}
results.loc['Nivel de Circularidad'] = pd.Series(hotel_results)#añadimos los datos como una nueva fila a la tabla resultados

#MOSTRAR LOS RESULTADOS
#display(results)
#results.to_excel("results.xlsx")

#FUNCIONES PARA GENERAR LOS GRÁFICOS
import plotly.graph_objects as go # or plotly.express as px
import plotly.graph_objects as go

#FUNCION PARA CREAR INDICADORES GAUGE DE NORMAL Y GUARDARLOS
def create_regular_indicator(indicador):
    i=1#ojo porque i se gasta y si no se pone al valor de la fila que queremos leer genera problemas
    if df_scale.Dir[indicador]==0:
        color_delta_inc="green"#color delta increasing
        color_delta_dec="red"#color delta decreasing
        n_min=df_scale.Min[indicador] #lee el valor mínimo de la hoja de escalas
        n_bajo=df_scale.Bajo[indicador] #lee el valor bajo de la hoja de escalas
        n_medio=df_scale.Medio[indicador]#lee el valor Medio de la hoja de escalas
        n_alto=df_scale.Alto[indicador] #lee el valor Alto de la hoja de escalas
        n_max=df_scale.Max[indicador] #lee el valor máximo de la hoja de escalas
        if df.iloc[i][indicador]<n_bajo: colores='red'#Conjunto de if para controlar el colorde la barra
        elif df.iloc[i][indicador]<n_medio: colores="yellow"
        elif df.iloc[i][indicador]<n_alto:colores="#85e043"
        else:colores="green"
    elif df_scale.Dir[indicador]==1:
        color_delta_inc="red"#color delta increasing
        color_delta_dec="green"#color delta decreasing
        n_min=df_scale.Max[indicador] #cambia la escala
        n_bajo=df_scale.Alto[indicador] 
        n_medio=df_scale.Medio[indicador]
        n_alto=df_scale.Bajo[indicador] 
        n_max=df_scale.Min[indicador] 
        if df.iloc[i][indicador]<n_bajo: colores="green"#Conjunto de if para controlar el colorde la barra
        elif df.iloc[i][indicador]<n_medio: colores="#85e043"
        elif df.iloc[i][indicador]<n_alto:colores="yellow"
        else:colores="red"

    
    fig = go.Figure(go.Indicator(
            mode = "gauge+number+delta",
            value = df.iloc[i][indicador],
        number = {'suffix': ''}, 
        domain = {'x': [0, 1], 'y': [0, 1]},
        delta = {'reference': df.iloc[i-1][indicador],'increasing': {'color': color_delta_inc},'decreasing': {'color': color_delta_dec}},
        gauge = {"axis": {'range': [n_min, n_max],
            "tickmode": "array",
            "tickvals": [n_min,n_bajo,n_medio, n_alto,n_max],
            "ticktext": [n_min,n_bajo,n_medio, n_alto,n_max]},
             'bar': {'color': colores},
                'threshold': {
            'line': {'color': "black", 'width': 6},
            'thickness': 0.75,
            'value': df.iloc[i-1][indicador]}},#divide la escala entre el máximo y mínimo y crea las divisiones de %
        
        title = {'text': diccionario_variables[indicador]}))
    filename='images/'+indicador+'.png'
    fig.write_image(filename)
    return filename
  


list_variables_percentage=['P1B','P2B','P3B','E3B','A1B','R2B','RA1B','RA2B', 'P2_3B']

#FUNCION PARA CREAR INDICADORES GAUGE DE PORCENTAJE Y GUARDARLOS
def create_percentage_indicator(indicador):
    i=1#para que coja el último año
    color_delta_inc="green"#color delta increasing
    color_delta_dec="red"#color delta decreasing
    if df.iloc[i][indicador]<df_scale.loc[indicador]["Bajo"]: colores='red'#Conjunto de if para controlar el colorde la barra
    elif df.iloc[i][indicador]<df_scale.loc[indicador]["Medio"]: colores="yellow"
    elif df.iloc[i][indicador]<df_scale.loc[indicador]["Alto"]:colores="#85e043"
    else:colores="green"
    fig = go.Figure(go.Indicator(
            mode = "gauge+number+delta",
            value = df.iloc[i][indicador],
        number = {'suffix': '%'}, 
        delta = {'reference': df.iloc[i-1][indicador],'increasing': {'color': color_delta_inc},'decreasing': {'color': color_delta_dec}},
        domain = {'x': [0, 1], 'y': [0, 1]},
        gauge = {"axis": {'range': [df_scale.loc[indicador]["Min"], df_scale.loc[indicador]["Max"]],
            "tickmode": "array",
            "tickvals": [df_scale.loc[indicador]["Min"],df_scale.loc[indicador]["Bajo"],df_scale.loc[indicador]["Medio"],
                         df_scale.loc[indicador]["Alto"],df_scale.loc[indicador]["Max"]],
            "ticktext": [str(df_scale.loc[indicador]["Min"])+'%',str(df_scale.loc[indicador]["Bajo"])+'%',str(df_scale.loc[indicador]["Medio"])+'%',str(df_scale.loc[indicador]["Alto"])+'%',str(df_scale.loc[indicador]["Max"])+'%']},
             'bar': {'color': colores},
                'threshold': {
            'line': {'color': "black", 'width': 6},
            'thickness': 0.75,
            'value': df.iloc[i-1][indicador]}},#divide la escala entre el máximo y mínimo y crea las divisiones de %
        title = {'text': diccionario_variables[indicador]}))
    filename='images/'+indicador+'.png'
    fig.write_image(filename)
    return filename

    

list_variables_certificate=['E9','E10','E11','E12','E2B', 'EI','E2B_val']

#FUNCION PARA CREAR INDICADORES GAUGE DE CERTIFICACIÓN
def create_certificate_indicator(indicador):
    i=1#para que coja el último año
    if df.iloc[i][indicador]<df_scale.loc[indicador]["Bajo"]: colores='red'#Conjunto de if para controlar el colorde la barra
    elif df.iloc[i][indicador]<df_scale.loc[indicador]["Medio"]: colores="yellow"
    elif df.iloc[i][indicador]<df_scale.loc[indicador]["Alto"]:colores="#85e043"
    else:colores="green"
    n_min=df_scale.Min[indicador] #lee el valor mínimo de la hoja de escalas
    n_bajo=df_scale.Bajo[indicador] #lee el valor bajo de la hoja de escalas
    n_medio=df_scale.Medio[indicador]#lee el valor Medio de la hoja de escalas
    n_alto=df_scale.Alto[indicador] #lee el valor Alto de la hoja de escalas
    n_max=df_scale.Max[indicador] #lee el valor máximo de la hoja de escalas
    fig = go.Figure(go.Indicator(
            mode = "gauge+number+delta",
            value = df.iloc[i][indicador], 
        delta = {'reference': df.iloc[i-1][indicador]},
        domain = {'x': [0, 1], 'y': [0, 1]},
        gauge = {"axis": {'range': [n_min, n_max],
            "tickmode": "array",
            "tickvals": [0,1/6,2/6,3/6,4/6,5/6,1],
            "ticktext": ['G','F','D','E','C','B', 'A']},
             'bar': {'color': colores},
                'threshold': {
            'line': {'color': "black", 'width': 6},
            'thickness': 0.75,
            'value': df.iloc[i-1][indicador]}},#divide la escala entre el máximo y mínimo y crea las divisiones de %
        
        title = {'text': diccionario_variables[indicador]}))
    filename='images/'+indicador+'.png'
    fig.write_image(filename)
    return filename


# FUNCIÓN PARA CREAR GRÁFICOS BAJO, MEDIO, ALTO, MUY ALTO
import numpy as np
def gauge_chart(var,value,min_value,max_value):
    plot_bgcolor = 'rgba(0, 0, 0, 0)'#"#1f2c56"#"#def"
    quadrant_colors = [plot_bgcolor,"#2bad4e", "#85e043", "#eff229","#f25829"  ] 
    quadrant_text = ["", "<b>Muy Alto</b>","<b>Alto</b>","<b>Medio</b>","<b>Bajo</b>"  ]
    n_quadrants = len(quadrant_colors) - 1
    current_value = value
    min_value = min_value
    max_value = max_value
    hand_length = np.sqrt(2) / 4
    hand_angle = np.pi * (1 - (max(min_value, min(max_value, current_value)) - min_value) / (max_value - min_value))

    fig = go.Figure(
        data=[
            go.Pie(
                values=[0.5] + (np.ones(n_quadrants) / 2 / n_quadrants).tolist(),
                rotation=90,
                hole=0.5,
                marker_colors=quadrant_colors,
                text=quadrant_text,
                textinfo="text",
                hoverinfo="skip",
            ),
        ],
        layout=go.Layout(
            showlegend=False,
            margin=dict(b=0,t=10,l=10,r=10),
            width=250,
            height=250,
            paper_bgcolor=plot_bgcolor,
            annotations=[
                go.layout.Annotation(
                    text=f"<b>{var} </b><br>{current_value} / {max_value} puntos",
                    x=0.5, xanchor="center", xref="paper",
                    y=0.25, yanchor="bottom", yref="paper",
                    showarrow=False,
                )
            ],
            shapes=[
                go.layout.Shape(
                    type="circle",
                    x0=0.48, x1=0.52,
                    y0=0.48, y1=0.52,
                    fillcolor="#333",
                    line_color="#333",
                ),
                go.layout.Shape(
                    type="line",
                    x0=0.5, x1=0.5 + hand_length * np.cos(hand_angle),
                    y0=0.5, y1=0.5 + hand_length * np.sin(hand_angle),
                    line=dict(color="#333", width=4)
                )
            ]
        )
    )
    filename=var+'.png'
    fig.write_image(filename)
    return filename


#GENERAR INFORME EN DOCX
from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL

dimensiones=['Planificación y Gestión Circular','Energía','Agua','Residuos','Desperdicio Alimentario']

document = Document()

document.add_heading('AUTOEVALUACIÓN DE CIRCULARIDAD')

#RESULTADO GENERAL
document.add_heading('NIVEL DE CIRCULARIDAD')
pic_name=gauge_chart('Nivel de Circularidad',results.loc['Nivel de Circularidad']['Puntuación asignada'],df_scale.loc['Nivel de Circularidad']['Min'],df_scale.loc['Nivel de Circularidad']['Max'])
document.add_picture(pic_name,width=1400000, height=1400000)


#DIMENSIONES
document.add_heading('DIMENSIONES DE CIRCULARIDAD')
#table with figures TO ADD DIMENSIONS
table = document.add_table(rows=2, cols=3)


x=0
i=0
for dimension in dimensiones:
#for dimension in dimensiones:
    pic_name=gauge_chart(dimension,results.loc[dimension]['Puntuación asignada'],df_scale.loc[dimension]['Min'],df_scale.loc[dimension]['Max'])
    cell = table.cell(i, x)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
    paragraph = cell.paragraphs[0]
    paragraph.alignment=WD_PARAGRAPH_ALIGNMENT.CENTER
    run = paragraph.add_run()
    run.add_text(dimension)
    run.add_break()
    run.add_picture(pic_name, width=1400000, height=1400000)  # size image
    if x==2: #si llegamos a la última columna
        i=1#pasamos a la siguiente fila
        x=0#reseteamos la columna
    else:
        x=x+1


#INDICADORES DE CADA DIMENSIÓN

for dimension in dimensiones:
    document.add_heading(dimension)
    table_dim = document.add_table(rows=1, cols=3)
    x=0
    i=0
    for n in range(0,len(results)):
        if results.iloc[n]['Dimensión']==dimension:
            var=results.iloc[n]['Indicador']
            pic_name=gauge_chart(results.iloc[n]['Subdimensión'],results.iloc[n]['Puntuación asignada'],0,df_scale.loc[var]['Puntos'])
            cell = table_dim.cell(i, x)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
            paragraph = cell.paragraphs[0]
            paragraph.alignment=WD_PARAGRAPH_ALIGNMENT.CENTER
            run = paragraph.add_run()
            run.add_text(results.iloc[n]['Subdimensión'])
            run.add_break()
            run.add_picture(pic_name, width=1400000, height=1400000)  # size image
            if x==2: #si llegamos a la última columna
                row=table_dim.add_row().cells#añadimos una fila más
                i=i+1#pasamos a la siguiente fila
                x=0#reseteamos la columna
            else:
                x=x+1


#INDICADORES DE CADA DIMENSIÓN con otro tipo de gauge chart

for dimension in dimensiones:
    document.add_heading(dimension)
    table_dim = document.add_table(rows=1, cols=2)
    x=0
    i=0
    for n in range(0,len(results)):
        if results.iloc[n]['Dimensión']==dimension:
            indicator=results.iloc[n]['Indicador']
            if indicator in list_variables_percentage:
                pic_name=create_percentage_indicator(indicator)
            elif indicator in list_variables_certificate:
                pic_name=create_certificate_indicator(indicator)
            else:
                pic_name=create_regular_indicator(indicator)
            cell = table_dim.cell(i, x)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
            paragraph = cell.paragraphs[0]
            paragraph.alignment=WD_PARAGRAPH_ALIGNMENT.CENTER
            run = paragraph.add_run()
            run.add_text(results.iloc[n]['Subdimensión'])
            run.add_break()
            run.add_picture(pic_name, width=1800000, height=1400000)  # size image
            if x==1: #si llegamos a la última columna
                row=table_dim.add_row().cells#añadimos una fila más
                i=i+1 #pasamos a la siguiente fila
                x=0 #reseteamos la columna
            else:
                x=x+1
          
            
var_habitacion_ocupada=['E62','E72','E82','E3B21','E3B22','E3B23','A1B2','A2B2',
                       'R12','R22','R32','R42','R52','R62'] 
document.add_heading('VARIABLES POR HABITACIÓN OCUPADA', level=2)


document.add_heading('VARIABLES POR HABITACIÓN OCUPADA', level=2)

#paragraph = document.add_paragraph('Lorem ipsum dolor sit amet.')
document.add_page_break()



#Each blank cell in a newly created table contains a single empty paragraph. This is just one of those things about the Word format. I suppose it gives a place to put the insertion mark (flashing vertical cursor) when you're using the Word application. A completely empty cell would have no place to "click" into.
#This requires that any code that adds content to a cell must treat the first paragraph differently. In short, you access the first paragraph as cell.paragraphs[0] and only create second and later paragraphs with cell.add_paragraph().

#table.style = 'LightShading-Accent1'
#document.add_picture('A1B.png', width=Inches(1.0))
document.save('test2.docx')


# ### 

# In[4]:


#pip install -U kaleido


# In[ ]:




